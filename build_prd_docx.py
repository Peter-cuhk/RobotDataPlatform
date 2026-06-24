from pathlib import Path
import re
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/peterxie/Desktop/data platform ")
SOURCE = ROOT / "PRD.md"
OUTPUT = ROOT / "Robot_Data_Studio_PRD_V1.0.docx"
ARCH_IMAGE = ROOT / ".prd_build" / "architecture.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B3A"
MUTED = "5F6B76"
LIGHT_BLUE = "E8F1F8"
LIGHT_GRAY = "F2F4F7"
LIGHT_GREEN = "EAF5EC"
LIGHT_GOLD = "FFF4DF"
WHITE = "FFFFFF"
RED = "9B1C1C"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, latin="Arial", east_asia="Arial Unicode MS", size=None,
                 color=None, bold=None, italic=None):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_paragraph_runs(paragraph, size=11, color=INK):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, size=9, color=MUTED)


def create_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167


def add_rich_text(paragraph, text, size=11, color=INK):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, latin="Menlo", east_asia="Arial Unicode MS", size=9.5,
                         color=DARK_BLUE)
            run.font.highlight_color = None
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + "  ")
    set_run_font(r, size=10.5, color=accent, bold=True)
    add_rich_text(p, text, size=10.5)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_summary_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, [2250, 7110])
    hdr = table.rows[0].cells
    hdr[0].text = "决策项"
    hdr[1].text = "已确认方案"
    for cell in hdr:
        set_cell_shading(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            for r in p.runs:
                set_run_font(r, size=10, color=INK, bold=True)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("产品形态", "完全开源、本地优先；浏览器 UI + 本地 Python API"),
        ("首版对象", "机械臂与具身操作 Episode 数据；单数据集不超过 10 GB"),
        ("核心格式", "LeRobot v2/v3、ACT/robomimic HDF5、UMI/Zarr profiles"),
        ("数据原则", "原始数据只读；所有清洗与转换使用可复现 Pipeline"),
        ("可视化", "Rerun 作为可替换 Viewer Adapter"),
        ("质量", "确定性规则为基础；VLM 作为可插拔辅助"),
        ("技术栈", "React/TypeScript + FastAPI/Python 模块化单体"),
        ("许可证", "建议 Apache-2.0；建立第三方许可证治理"),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for idx, cell in enumerate(cells):
            if idx == 0:
                set_cell_shading(cell, "F8FAFC")
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run_font(r, size=9.5, color=INK, bold=(idx == 0))


def create_architecture_image(path):
    path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 1500, 760
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_paths = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
    ]
    fp = next((p for p in font_paths if Path(p).exists()), None)
    title = ImageFont.truetype(fp, 42) if fp else ImageFont.load_default()
    heading = ImageFont.truetype(fp, 30) if fp else ImageFont.load_default()
    body = ImageFont.truetype(fp, 23) if fp else ImageFont.load_default()
    small = ImageFont.truetype(fp, 20) if fp else ImageFont.load_default()

    draw.text((70, 35), "Robot Data Studio — Local-first Architecture", fill="#172B3A", font=title)

    def box(x, y, w, h, fill, outline, heading_text, lines):
        draw.rounded_rectangle((x, y, x + w, y + h), radius=18, fill=fill,
                               outline=outline, width=4)
        draw.text((x + 24, y + 20), heading_text, fill=outline, font=heading)
        yy = y + 72
        for line in lines:
            draw.text((x + 24, yy), line, fill="#314555", font=body)
            yy += 38

    box(70, 125, 390, 190, "#F8FAFC", "#607D94", "只读原始数据",
        ["LeRobot", "HDF5 profiles", "UMI / Zarr"])
    box(555, 125, 390, 190, "#EEF7FF", "#2E74B5", "Canonical Episode View",
        ["Reader Adapters", "按字段 / Episode / 时间窗口读取", "不强制转存"])
    box(1040, 125, 390, 190, "#EFF8EF", "#4F8358", "Cleaning Pipeline",
        ["Quality Gate", "Trim / Resample", "Coordinate Transform"])
    draw.line((460, 220, 555, 220), fill="#6B7E8D", width=7)
    draw.polygon([(555, 220), (530, 206), (530, 234)], fill="#6B7E8D")
    draw.line((945, 220, 1040, 220), fill="#6B7E8D", width=7)
    draw.polygon([(1040, 220), (1015, 206), (1015, 234)], fill="#6B7E8D")

    box(70, 410, 390, 175, "#F3F0FF", "#7457A6", "Rerun Viewer",
        ["视频 / 曲线 / 3D", "统一时间轴", "Finding 跳转"])
    box(555, 410, 390, 175, "#FFF5E8", "#B56A16", "Quality Engine",
        ["确定性规则", "可插拔 VLM", "可解释 Findings"])
    box(1040, 410, 390, 175, "#F1F5F9", "#466B8E", "Writer Adapters",
        ["LeRobot", "ACT / robomimic HDF5", "UMI / Zarr"])

    draw.rounded_rectangle((220, 650, 1280, 720), radius=14, fill="#F4F6F9",
                           outline="#A8B5C0", width=2)
    draw.text((250, 670),
              "本地项目只保存：SQLite 元数据 · Pipeline JSON · 缩略图/统计缓存 · 质量与导出报告",
              fill="#4D5C68", font=small)
    image.save(path)


def add_cover(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.text = "ROBOT DATA STUDIO  |  PRODUCT REQUIREMENTS"
    for run in header.runs:
        set_run_font(run, size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(40)
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(10)
    run = kicker.add_run("OPEN-SOURCE PRODUCT REQUIREMENTS DOCUMENT")
    set_run_font(run, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Robot Data Studio")
    set_run_font(run, size=28, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(26)
    run = subtitle.add_run("本地优先的机器人数据可视化、质量审查、清洗、坐标与格式转换平台")
    set_run_font(run, size=15, color=DARK_BLUE)

    add_callout(
        doc,
        "产品愿景",
        "把分散的回放工具、数据质量脚本、Schema 转换器和坐标处理代码，整合成一个可信、可解释、可复现的开源机器人数据工作台。",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )

    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(meta, [2100, 7260])
    metadata = [
        ("版本", "V1.0 Draft"),
        ("日期", "2026-06-23"),
        ("产品形态", "完全开源 / Local-first / Browser UI + Local API"),
        ("MVP 范围", "Manipulation Episodes；LeRobot、HDF5、UMI/Zarr；≤10 GB"),
        ("状态", "产品范围与核心架构已确认"),
    ]
    for row, (label, value) in zip(meta.rows, metadata):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run_font(r, size=10, color=INK, bold=(idx == 0))

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_front_matter(doc):
    h = doc.add_heading("执行摘要", level=1)
    h.paragraph_format.space_before = Pt(0)
    p = doc.add_paragraph()
    add_rich_text(
        p,
        "本 PRD 定义一个完全开源、本地优先的机器人数据清洗平台。首版聚焦机械臂与具身操作 Episode 数据，以统一语义层连接回放、质量审查、非破坏式清洗、坐标变换和格式导出。",
    )
    add_summary_table(doc)
    doc.add_paragraph()
    create_architecture_image(ARCH_IMAGE)
    picture = doc.add_picture(str(ARCH_IMAGE), width=Inches(6.35))
    inline = picture._inline
    doc_pr = inline.docPr
    doc_pr.set("name", "Robot Data Studio architecture")
    doc_pr.set("descr", "Local-first architecture showing read-only source data, canonical episode view, cleaning pipeline, Rerun viewer, quality engine, and writer adapters.")
    cap = doc.add_paragraph("图 1  推荐的本地优先模块化架构")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    for run in cap.runs:
        set_run_font(run, size=9, color=MUTED, italic=True)

    doc.add_heading("阅读指引", level=2)
    p = doc.add_paragraph()
    add_rich_text(p, "第 1–5 节解释产品动机与范围；第 6–8 节定义用户工作流、功能与 Canonical Schema；第 9–11 节给出技术选型、插件和开源复用策略；第 12–18 节覆盖非功能需求、验收、路线图、风险和参考资料。")


def parse_and_add_markdown(doc, lines):
    in_code = False
    code_lines = []
    skip_front = True
    numbered_pattern = re.compile(r"^\d+\.\s+")
    section_count = 0
    active_numbering_id = None

    for raw in lines:
        line = raw.rstrip()
        if skip_front:
            if line.startswith("## 1. "):
                skip_front = False
            else:
                continue

        if line.startswith("```"):
            if in_code:
                table = doc.add_table(rows=1, cols=1)
                set_table_widths(table, [9360])
                cell = table.cell(0, 0)
                set_cell_shading(cell, "F5F7F9")
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                for idx, code_line in enumerate(code_lines):
                    if idx:
                        p.add_run("\n")
                    run = p.add_run(code_line)
                    set_run_font(run, latin="Menlo", east_asia="Arial Unicode MS",
                                 size=8.5, color="263746")
                doc.add_paragraph().paragraph_format.space_after = Pt(1)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue

        if not line:
            active_numbering_id = None
            continue
        if line.startswith("## "):
            active_numbering_id = None
            section_count += 1
            if section_count > 1:
                # Major sections start cleanly when near a page boundary, while Word may
                # still keep compact adjacent sections together.
                pass
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            active_numbering_id = None
            doc.add_heading(line[4:], level=2)
        elif line.startswith("#### "):
            active_numbering_id = None
            doc.add_heading(line[5:], level=3)
        elif line.startswith("- "):
            active_numbering_id = None
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, line[2:])
        elif numbered_pattern.match(line):
            if active_numbering_id is None:
                active_numbering_id = create_decimal_numbering(doc)
            p = doc.add_paragraph()
            apply_numbering(p, active_numbering_id)
            add_rich_text(p, numbered_pattern.sub("", line))
        else:
            active_numbering_id = None
            p = doc.add_paragraph()
            add_rich_text(p, line)


def add_technical_decision_table(doc):
    doc.add_heading("附录 A：技术选型决策矩阵", level=1)
    table = doc.add_table(rows=1, cols=4)
    set_table_widths(table, [1700, 2300, 2860, 2500])
    headers = ["领域", "选择", "采用原因", "明确不选/注意事项"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_margins(cell, top=35, bottom=35, start=100, end=100)
        cell.text = text
        set_cell_shading(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            for r in p.runs:
                set_run_font(r, size=9, color=INK, bold=True)
    rows = [
        ("产品架构", "模块化单体", "本地部署简单；接口仍可替换", "MVP 不拆微服务"),
        ("前端", "React + TypeScript + Vite", "适合复杂状态、插件 UI 和长期维护", "不以 Streamlit/Gradio 作为正式壳"),
        ("本地 API", "FastAPI + Pydantic", "Python 数据生态与类型化 API", "后台重任务不能只用 in-process BackgroundTasks"),
        ("回放", "Rerun Web Viewer", "同步时间轴、视频、曲线和 3D 能力成熟", "通过 Adapter 隔离，不绑定核心 Schema"),
        ("Pipeline UI", "React Flow", "节点编辑成熟，MIT", "Pipeline IR 归平台所有"),
        ("元数据", "SQLite", "零运维、本地事务可靠", "不存视频和大型数组"),
        ("数据访问", "h5py / zarr / PyArrow / PyAV", "按格式使用官方主流库", "禁止全量载入内存"),
        ("后台任务", "ProcessPoolExecutor", "首版足够、可取消和隔离 CPU 工作", "保留 TaskExecutor 接口以升级 Ray/Celery"),
        ("开源许可证", "Apache-2.0", "宽松且包含专利授权条款", "持续扫描第三方许可证"),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            cells[idx].text = text
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cells[idx], top=35, bottom=35, start=100, end=100)
            for p in cells[idx].paragraphs:
                for r in p.runs:
                    set_run_font(r, size=8.2, color=INK, bold=(idx == 0))


def add_footer_note(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Product Requirements Document —")
    set_run_font(run, size=9, color=MUTED, italic=True)


def build():
    text = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_front_matter(doc)
    parse_and_add_markdown(doc, text)
    add_technical_decision_table(doc)

    props = doc.core_properties
    props.title = "Robot Data Studio 产品需求文档（PRD）"
    props.subject = "Local-first open-source robot data cleaning and conversion platform"
    props.author = "Robot Data Studio Project"
    props.keywords = "robotics, data quality, HDF5, LeRobot, UMI, Rerun, PRD"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
